
from pathlib import Path
import tempfile,json
from fastapi import FastAPI,UploadFile,File,Form,HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from io import BytesIO
from docx import Document
from analysis_engine import analyse_file,analyse_text,analyse_photo_placeholder
app=FastAPI(title="VinQu Backend")
app.add_middleware(CORSMiddleware,allow_origins=["*"],allow_methods=["*"],allow_headers=["*"])
@app.get("/")
def root(): return {"status":"ok","message":"VinQu backend is running"}
@app.post("/analyse/text")
async def analyse_text_endpoint(text:str=Form(...),file_name:str=Form("")): return analyse_text(text,file_name)
@app.post("/analyse/file")
async def analyse_file_endpoint(file:UploadFile=File(...),referenced_docs:list[UploadFile]=File(default=[])):
 suffix=Path(file.filename or "").suffix
 if not suffix: raise HTTPException(status_code=400,detail="Uploaded file has no file extension.")
 with tempfile.NamedTemporaryFile(delete=False,suffix=suffix) as tmp:
  tmp.write(await file.read()); temp_path=tmp.name
 try:
  return analyse_file(temp_path,[r.filename for r in referenced_docs if r.filename])
 except Exception as exc: raise HTTPException(status_code=400,detail=str(exc))
@app.post("/analyse/photo")
async def analyse_photo_endpoint(photo:UploadFile=File(...),inspection_area:str=Form(""),inspection_item:str=Form(""),inspector_notes:str=Form(""),document_context:str=Form("{}")):
 return analyse_photo_placeholder(photo.filename or "",inspection_area,inspection_item,inspector_notes,document_context)
@app.post("/report/docx")
async def report_docx(payload:str=Form(...)):
 try: data=json.loads(payload)
 except Exception: raise HTTPException(status_code=400,detail="Invalid report payload.")
 doc=Document(); doc.add_heading("VinQu Consolidated Inspection Report",0); ins=data.get("instruction") or {}; doc.add_heading("Inspection Summary",level=1)
 for key in ["document_type","project_reference","inspection_date_window","inspection_factory_location","equipment","scope_of_work","reasoned_summary"]: doc.add_paragraph(f"{key.replace('_',' ').title()}: {ins.get(key,'-')}")
 doc.add_heading("Inspection Activities / Tests",level=1)
 for i,item in enumerate(ins.get("tests_or_checks") or [],1): doc.add_paragraph(f"{i}. {item}")
 doc.add_heading("Findings",level=1)
 for i,f in enumerate(data.get("findings") or [],1):
  doc.add_heading(f"Finding {i}",level=2); doc.add_paragraph(f"Inspection item: {f.get('inspectionItem','-')}"); doc.add_paragraph(f"Area / component: {f.get('inspectionArea','-')}"); doc.add_paragraph(f"Inspector notes: {f.get('inspectorNotes','-')}"); r=f.get("result") or {}; doc.add_paragraph(f"Status: {r.get('status_label','-')}"); doc.add_paragraph(f"Finding: {r.get('suggested_finding','-')}"); doc.add_paragraph(f"Reasoning: {r.get('reasoning','-')}"); doc.add_paragraph(f"Standards reasoning: {r.get('standard_reasoning','-')}")
 bio=BytesIO(); doc.save(bio); bio.seek(0)
 return StreamingResponse(bio,media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",headers={"Content-Disposition":"attachment; filename=VinQu_Consolidated_Inspection_Report.docx"})
